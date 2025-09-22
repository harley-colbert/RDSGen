from __future__ import annotations
from pathlib import Path
from docxtpl import DocxTemplate
from jinja2 import Environment, Undefined
from ..domain.models import Inputs, Computation

class WordGenerator:
    def __init__(self, template_path: str):
        self.template_path = template_path

    def generate(self, out_path: Path, inputs: Inputs, comp: Computation):
        if not self.template_path:
            from docx import Document
            doc = Document()
            doc.add_heading('RDSGen Quote', level=1)
            doc.add_paragraph(f'Base Price: {inputs.base_price:,.2f}')
            doc.add_paragraph(f'Options Price: {comp.options_price_total:,.2f}')
            doc.add_paragraph(f'Total: {comp.total_price:,.2f}')
            doc.save(out_path)
            return
        env = Environment(undefined=Undefined)
        tpl = DocxTemplate(self.template_path)
        ctx = {
            "base_price": inputs.base_price,
            "options_price": comp.options_price_total,
            "total_price": comp.total_price,
            "margin": inputs.margin,
            "options": [{"name": k, "qty": comp.options_qty.get(k, 1), "extended": v} for k, v in comp.options_breakdown.items()],
            "inputs": inputs.model_dump(),
        }
        tpl.render(ctx, environment=env)
        tpl.save(out_path)
